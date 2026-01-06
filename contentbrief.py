# contentbrief.py — simple two-column layout, st.json rendering (no CSS, no st.code)

from __future__ import annotations

import io
import json
import re
from typing import Any, Dict, List

import requests
import streamlit as st
from docx import Document

# Your existing shared utilities (unchanged)
from shared import (
    serpapi_google_search,
    extract_serp_info,                # -> headings, fs_single, related_questions, ai_overview_raw, related_searches
    extract_semantic_terms_with_llm,  # -> (terms: list[str], clusters: list[{"name","items"}])
    openai_chat,                      # OpenAI wrapper
)

# ------------------------------
# Small helpers
# ------------------------------

def _lines(text: str) -> list[str]:
    text = (text or "").strip()
    return [ln for ln in re.split(r"\r?\n", text) if ln.strip()]

def _add_unique(acc: list[str], s: str):
    s = (s or "").strip()
    if s and s not in acc:
        acc.append(s)

# ------------------------------
# Featured Snippet collector (ALL variants)
# ------------------------------

FS_KEYS_PARA = ("snippet", "answer", "title")
FS_KEYS_LIST = ("list",)
FS_KEYS_TABLE = ("table",)

def _fs_from_obj(obj: Dict[str, Any]) -> list[str]:
    out: list[str] = []

    for k in FS_KEYS_PARA:
        v = obj.get(k)
        if isinstance(v, str) and v.strip():
            _add_unique(out, v)

    hl = obj.get("snippet_highlighted_words")
    if isinstance(hl, list) and any(isinstance(x, str) for x in hl):
        _add_unique(out, " ".join(x for x in hl if isinstance(x, str) and x.strip()))

    for k in FS_KEYS_LIST:
        v = obj.get(k)
        if isinstance(v, list) and v:
            items = [str(i).strip() for i in v if str(i).strip()]
            if items:
                _add_unique(out, "\n".join(items))

    for k in FS_KEYS_TABLE:
        tbl = obj.get(k)
        if tbl:
            if isinstance(tbl, dict) and isinstance(tbl.get("rows"), list):
                rows = [" | ".join(str(c) for c in r if c is not None) for r in tbl["rows"]]
                rows = [r for r in rows if r.strip()]
                if rows:
                    _add_unique(out, "\n".join(rows))
            elif isinstance(tbl, list):
                rows = []
                for r in tbl:
                    if isinstance(r, (list, tuple)):
                        rows.append(" | ".join(str(c) for c in r))
                    else:
                        rows.append(str(r))
                rows = [r for r in rows if r.strip()]
                if rows:
                    _add_unique(out, "\n".join(rows))

    return out

def collect_all_featured_snippets(serp_json: Dict[str, Any]) -> list[str]:
    out: list[str] = []

    def walk(x):
        if isinstance(x, dict):
            if "featured_snippet" in x and isinstance(x["featured_snippet"], dict):
                for val in _fs_from_obj(x["featured_snippet"]):
                    _add_unique(out, val)
            if "answer_box" in x and isinstance(x["answer_box"], dict):
                for val in _fs_from_obj(x["answer_box"]):
                    _add_unique(out, val)
            for v in x.values():
                walk(v)
        elif isinstance(x, list):
            for v in x:
                walk(v)

    walk(serp_json or {})
    seen, clean = set(), []
    for s in out:
        if s not in seen:
            seen.add(s)
            clean.append(s)
    return clean[:20]

# ------------------------------
# AI Overview / SGE collector (ALL)
# ------------------------------

def _ai_blocks_from_text_blocks(text_blocks) -> list[dict]:
    blocks: list[dict] = []
    for b in (text_blocks or []):
        btype = (b.get("type") or "").lower()
        if btype == "paragraph":
            title = (b.get("title") or "").strip()
            if title:
                blocks.append({"type": "paragraph", "text": title})
            snip = (b.get("snippet") or b.get("text") or "").strip()
            if snip:
                blocks.append({"type": "paragraph", "text": snip})
        elif btype == "list":
            items = []
            for li in (b.get("list") or []):
                if isinstance(li, dict):
                    items.append((li.get("title") or li.get("snippet") or "").strip())
                else:
                    items.append(str(li).strip())
            items = [i for i in items if i]
            if items:
                blocks.append({"type": "list", "items": items})
    return blocks

def collect_all_ai_overview_blocks(ai_overview_obj_or_text) -> list[dict]:
    if isinstance(ai_overview_obj_or_text, str):
        txt = ai_overview_obj_or_text.strip()
        return [{"type": "paragraph", "text": txt}] if txt else []

    out: list[dict] = []

    def add_para(s: str):
        s = (s or "").strip()
        if s:
            out.append({"type": "paragraph", "text": s})

    def add_list(seq):
        items = [str(i).strip() for i in seq if str(i).strip()]
        if items:
            out.append({"type": "list", "items": items})

    def walk(v):
        if isinstance(v, dict):
            if isinstance(v.get("text_blocks"), list):
                out.extend(_ai_blocks_from_text_blocks(v["text_blocks"]))

            for hk in ("heading", "title", "label"):
                hv = v.get(hk)
                if isinstance(hv, str) and hv.strip():
                    add_para(hv)

            for k in ("text", "snippet", "answer"):
                tv = v.get(k)
                if isinstance(tv, str) and tv.strip():
                    add_para(tv)

            if isinstance(v.get("list"), list):
                add_list(v["list"])

            for k in ("sections", "entries", "cards", "chips", "items", "content", "results", "blocks", "children"):
                if k in v and isinstance(v[k], (list, dict)):
                    walk(v[k])

            for vv in v.values():
                if isinstance(vv, (list, dict)):
                    walk(vv)

        elif isinstance(v, list):
            for vv in v:
                walk(vv)

    walk(ai_overview_obj_or_text or {})

    seen, clean = set(), []
    for b in out:
        key = json.dumps(b, ensure_ascii=False, sort_keys=True)
        if key not in seen:
            seen.add(key)
            clean.append(b)
    return clean[:100]

# ------------------------------
# Term classification (weight + intent)
# ------------------------------

def classify_terms_with_llm(api_key: str, terms: list[str], model: str) -> list[dict]:
    prompt = f"""
Assign each term a weight (1–5) and an intent: informational, commercial, or transactional.
Return ONLY a JSON array like:
[{{"term":"best vpn for firestick","weight":5,"intent":"commercial"}}, ...]
Terms:
{json.dumps(terms, ensure_ascii=False)}
""".strip()

    raw = openai_chat(
        api_key,
        [
            {"role": "system", "content": "You are a precise SEO assistant. Output JSON only."},
            {"role": "user", "content": prompt},
        ],
        temperature=0,
        model=model,
    )

    try:
        data = json.loads(raw)
    except Exception:
        m = re.search(r"\[.*\]", raw, re.S)
        data = json.loads(m.group(0)) if m else []

    out = []
    for item in (data or []):
        try:
            term = str(item.get("term", "")).lower().strip()
            weight = int(item.get("weight", 3))
            weight = max(1, min(5, weight))
            intent = str(item.get("intent", "—")).lower().strip()
            if intent not in ("informational", "commercial", "transactional"):
                intent = "—"
            if term:
                out.append({"term": term, "weight": weight, "intent": intent})
        except Exception:
            continue

    if not out:
        out = [{"term": t, "weight": 3, "intent": "—"} for t in terms]
    return out

# ------------------------------
# DOCX builders (kept as before)
# ------------------------------

def _docx_serp_signals(keyword, headings, featured_snippets, ai_blocks, related_questions, related_searches) -> io.BytesIO:
    doc = Document()
    doc.add_heading(f"SERP Signals for: {keyword}", level=1)

    doc.add_heading("Top Result Headings", level=2)
    for h in (headings or []):
        doc.add_paragraph(h, style="List Bullet")

    doc.add_heading("Featured Snippet(s)", level=2)
    if featured_snippets:
        for fs in featured_snippets:
            for ln in _lines(fs):
                doc.add_paragraph(ln)
            doc.add_paragraph("")
    else:
        doc.add_paragraph("—")

    doc.add_heading("AI Overview / SGE", level=2)
    if ai_blocks:
        for b in ai_blocks:
            if b["type"] == "paragraph":
                for ln in _lines(b["text"]):
                    doc.add_paragraph(ln)
            elif b["type"] == "list":
                for it in b["items"]:
                    doc.add_paragraph(it, style="List Bullet")
            doc.add_paragraph("")
    else:
        doc.add_paragraph("—")

    doc.add_heading("People Also Ask", level=2)
    for q in (related_questions or []):
        doc.add_paragraph(q, style="List Bullet")

    doc.add_heading("Related Searches", level=2)
    for rs in (related_searches or []):
        doc.add_paragraph(rs, style="List Bullet")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def _docx_terms_clusters(keyword: str, classified_terms: list[dict], clusters: list[dict]) -> io.BytesIO:
    doc = Document()
    doc.add_heading(f"Semantic Terms & Clusters for: {keyword}", level=1)

    doc.add_heading("Terms", level=2)
    if classified_terms:
        for item in classified_terms:
            doc.add_paragraph(
                f'{item["term"]}  —  w:{item["weight"]}  ·  {item["intent"]}',
                style="List Bullet"
            )
    else:
        doc.add_paragraph("—")

    doc.add_heading("Clusters", level=2)
    for c in (clusters or []):
        name = (c.get("name") or "").strip()
        items = (c.get("items") or []) or []
        if name:
            doc.add_heading(name, level=3)
        if items:
            for it in items:
                doc.add_paragraph(str(it), style="List Bullet")
        else:
            doc.add_paragraph("—")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def _docx_brief(keyword: str, brief_text: str) -> io.BytesIO:
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(f"Content Brief: {keyword}", level=1)
    for line in (brief_text or "").splitlines():
        if line.strip().startswith("# "):
            doc.add_heading(line.replace("# ", "").strip(), level=2)
        elif line.strip().startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=3)
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ------------------------------
# Brief generation
# ------------------------------

def generate_content_brief(
    api_key: str,
    keyword: str,
    headings,
    featured_snippets: list[str],
    ai_blocks: list[dict],
    related_questions,
    classified_terms: list[dict],
    clusters: list[dict] | None = None,
    model: str = "gpt-4o-mini",
    temperature: float = 0.25,
) -> str:
    top_headings = [str(h).strip() for h in (headings or []) if str(h).strip()]
    fs_text = (featured_snippets[0] if featured_snippets else None)
    # stringify AO blocks for the prompt
    def _stringify_ai_overview_for_prompt(ai_blocks: list[dict]) -> str:
        parts = []
        for b in (ai_blocks or []):
            t = (b.get("type") or "").lower()
            if t == "paragraph":
                txt = (b.get("text") or "").strip()
                if txt:
                    parts.append(txt)
            elif t == "list":
                items = [str(i).strip() for i in (b.get("items") or []) if str(i).strip()]
                if items:
                    parts.append(", ".join(items))
        s = " ".join(parts).strip()
        return s[:3000]
    ao_text = _stringify_ai_overview_for_prompt(ai_blocks or [])
    paa = [str(q).strip() for q in (related_questions or []) if str(q).strip()]
    semantic_terms = [t.get("term", "").strip() for t in (classified_terms or []) if str(t.get("term","")).strip()]

    clusters = clusters or []
    cluster_lines = []
    for c in clusters:
        name = (c.get("name") or "").strip()
        items = [str(i).strip() for i in (c.get("items") or []) if str(i).strip()]
        if name and items:
            cluster_lines.append(f"- {name}: {', '.join(items)}")
        elif items:
            cluster_lines.append(f"- {', '.join(items)}")
    clusters_text = "\n".join(cluster_lines) if cluster_lines else "None"

    prompt = f"""
Based on these Google search insights for the keyword "{keyword}":

- Top page titles/headings: {', '.join(top_headings) if top_headings else 'None'}
- Featured snippet text: {fs_text or 'None'}
- AI Overview (Google SGE): {ao_text or 'None'}
- Related questions (PAA): {', '.join(paa) if paa else 'None'}
- Semantic terms to weave in: {', '.join(semantic_terms) if semantic_terms else 'None'}
- Semantic clusters:
{clusters_text}

Write a detailed SEO content overview and a structured content brief outlining:
- Main topics to cover
- Suggested article sections/headings
- For each section, specify the content type (paragraphs, bullet list, numbered steps, table, CTA box, code/example, FAQ, etc.)
- Where appropriate, indicate which semantic terms should naturally appear in that section (do NOT keyword-stuff; keep it natural).
- A concise FAQ list

Optimize for search intent and readability. Keep instructions concrete and implementation-ready.
""".strip()

    try:
        resp = openai_chat(
            api_key=api_key,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            model=model,
        )
        return resp.strip() if isinstance(resp, str) else json.dumps(resp, ensure_ascii=False)
    except Exception as e:
        return f"(brief generation failed) {e}"


# ------------------------------
# UI (left/right, st.json outputs)
# ------------------------------

def render(OPENAI_API_KEY: str, SERPAPI_KEY: str, gl: str, hl: str, OPENAI_MODEL: str = "gpt-4o-mini", show_header: bool = False):
    KEY = "cb_state"
    if KEY not in st.session_state:
        st.session_state[KEY] = {}
    S = st.session_state[KEY]

    if show_header:
        st.subheader("Content Brief")
    st.write("Generate a structured SEO content brief from SERP signals, AI Overviews, and semantic terms.")

    kw = st.text_input("Primary keyword", value=S.get("keyword", "best vpn for firestick"), key="cb_kw")
    num = st.number_input("Results to fetch (1–100)", min_value=1, max_value=100, value=int(S.get("num", 20)), step=1)
    run = st.button("Run Content Brief", type="primary")

    if run and (not SERPAPI_KEY or not OPENAI_API_KEY):
        st.error("Missing SERPAPI_KEY or OPENAI_API_KEY.")

    if run and SERPAPI_KEY and OPENAI_API_KEY:
        try:
            with st.status("Fetching SERP data from SerpAPI…", expanded=False):
                serp_json = serpapi_google_search(kw, int(num), gl, hl, SERPAPI_KEY)
                st.write("✅ SERP fetched")

            with st.status("Extracting SERP signals…", expanded=False):
                headings, fs_single, related_questions, ai_overview_raw, related_searches = extract_serp_info(serp_json)
                featured_snippets = collect_all_featured_snippets(serp_json)
                if fs_single and fs_single.strip():
                    _add_unique(featured_snippets, fs_single.strip())
                ai_blocks = collect_all_ai_overview_blocks(ai_overview_raw)
                st.write("✅ Signals extracted")

            with st.status("Deriving semantic terms & clusters via LLM…", expanded=False):
                terms, clusters = extract_semantic_terms_with_llm(
                    headings=headings,
                    snippet="\n\n".join(featured_snippets[:2]) if featured_snippets else "",
                    related_questions=related_questions,
                    ai_overview=ai_overview_raw,
                    related_searches=related_searches,
                    keyword=kw,
                    max_terms=25,
                    api_key=OPENAI_API_KEY,
                    model=OPENAI_MODEL,
                    temperature=0.2,
                )
                classified_terms = classify_terms_with_llm(OPENAI_API_KEY, terms, model=OPENAI_MODEL)
                st.write("✅ Terms classified")

            with st.status("Generating content brief…", expanded=False):
                brief_text = generate_content_brief(
                    OPENAI_API_KEY,
                    kw,
                    headings,
                    featured_snippets,
                    ai_blocks,
                    related_questions,
                    classified_terms,
                    clusters=clusters,
                    model=OPENAI_MODEL,
                )
                st.write("✅ Brief generated")

            serp_docx = _docx_serp_signals(kw, headings, featured_snippets, ai_blocks, related_questions, related_searches)
            terms_docx = _docx_terms_clusters(kw, classified_terms, clusters)
            brief_docx = _docx_brief(kw, brief_text)

            S.update({
                "keyword": kw,
                "num": int(num),
                "serp_json": serp_json,
                "headings": headings,
                "featured_snippets": featured_snippets,
                "ai_blocks": ai_blocks,
                "related_questions": related_questions,
                "related_searches": related_searches,
                "terms": terms,                          # strings list
                "classified_terms": classified_terms,     # dicts list (kept for downloads)
                "clusters": clusters,
                "brief_text": brief_text,
                "serp_docx_bytes": serp_docx.getvalue(),
                "terms_docx_bytes": terms_docx.getvalue(),
                "brief_docx_bytes": brief_docx.getvalue(),
                "ready": True,
            })
            st.success("Pipeline completed. Downloads are now persistent until you run again or change inputs.")

        except requests.Timeout:
            st.error("SerpAPI timed out. Try fewer results.")
        except requests.HTTPError as e:
            st.error(f"SerpAPI HTTP error: {e}")
        except Exception as e:
            st.error(f"Unexpected error: {e}")

    if S.get("ready"):
        st.header("SERP Signals")
        left, right = st.columns(2, gap="large")

        with left:
            st.subheader("Top result headings")
            st.json(S.get("headings") or [])

            st.subheader("People Also Ask")
            st.json(S.get("related_questions") or [])

            st.subheader("Related searches")
            st.json(S.get("related_searches") or [])

            st.download_button(
                "Download SERP Signals (.docx)",
                data=S.get("serp_docx_bytes", b""),
                file_name=f"serp_signals_{S['keyword'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_serp_docx",
            )

        with right:
            st.subheader("Featured Snippet")
            st.json(S.get("featured_snippets") or [])

            st.subheader("AI Overview (blocks)")
            st.json(S.get("ai_blocks") or [])

        st.header("Semantic Terms & Clusters")
        tleft, tright = st.columns(2, gap="large")
        with tleft:
            st.subheader("Terms")
            # show the raw terms list (matches your screenshot)
            st.json(S.get("terms") or [])
        with tright:
            st.subheader("Clusters")
            st.json(S.get("clusters") or [])

        st.download_button(
            "Download Terms & Clusters (.docx)",
            data=S.get("terms_docx_bytes", b""),
            file_name=f"semantic_terms_clusters_{S['keyword'].replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_terms_docx",
        )

        st.header("Content Brief")
        st.text(S.get("brief_text") or "")
        
        st.download_button(
            "Download Content Brief (.docx)",
            data=S.get("brief_docx_bytes", b""),
            file_name=f"content_brief_{S['keyword'].replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_brief_docx",
        )

    else:
        st.info("Run the pipeline to view SERP signals, AI Overview blocks, semantic terms/clusters, and the generated brief.")
